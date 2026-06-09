"""Build a basic local text index for PDF, DOCX, and LaTeX literature.

Dependencies:
    python -m pip install pypdf python-docx

This script performs text extraction only. Scanned PDFs require a separate OCR
workflow and are recorded with an explanatory warning when no text is found.
"""

from __future__ import annotations

import argparse
import hashlib
import json
import re
from datetime import datetime, timezone
from pathlib import Path
from typing import Callable

from common import DEFAULT_INDEX_PATH, MISSING, PROJECT_ROOT, normalize_space

SUPPORTED_SUFFIXES = {".pdf", ".docx", ".tex"}


def read_pdf(path: Path) -> str:
    """Extract text from a text-based PDF."""
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    return "\n\n".join(page.extract_text() or "" for page in reader.pages)


def read_docx(path: Path) -> str:
    """Extract paragraph and table text from a DOCX file."""
    from docx import Document

    document = Document(str(path))
    blocks = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            blocks.append("\t".join(cell.text for cell in row.cells))
    return "\n".join(blocks)


def read_tex(path: Path) -> str:
    """Read LaTeX source and remove common comments and markup."""
    source = path.read_text(encoding="utf-8", errors="replace")
    source = re.sub(r"(?<!\\)%.*", "", source)
    source = re.sub(
        r"\\(?:section|subsection|subsubsection|paragraph)\*?\{([^{}]*)\}",
        r"\n\1\n",
        source,
    )
    source = re.sub(r"\\(?:cite|ref|label|includegraphics)\*?(?:\[[^\]]*\])?\{[^{}]*\}", " ", source)
    source = re.sub(r"\\[a-zA-Z@]+\*?(?:\[[^\]]*\])?", " ", source)
    source = source.replace("{", " ").replace("}", " ")
    return source


READERS: dict[str, Callable[[Path], str]] = {
    ".pdf": read_pdf,
    ".docx": read_docx,
    ".tex": read_tex,
}


def infer_title(text: str, path: Path) -> str:
    """Infer a cautious title candidate from explicit markup or first lines."""
    if path.suffix.lower() == ".tex":
        source = path.read_text(encoding="utf-8", errors="replace")
        match = re.search(r"\\title\{(.+?)\}", source, flags=re.DOTALL)
        if match:
            return normalize_space(match.group(1))
    lines = [normalize_space(line) for line in text.splitlines()]
    candidates = [line for line in lines[:40] if 5 <= len(line) <= 240]
    return candidates[0] if candidates else path.stem


def infer_abstract(text: str, path: Path) -> str:
    """Infer an abstract candidate without claiming semantic certainty."""
    if path.suffix.lower() == ".tex":
        source = path.read_text(encoding="utf-8", errors="replace")
        match = re.search(
            r"\\begin\{abstract\}(.*?)\\end\{abstract\}",
            source,
            flags=re.DOTALL | re.IGNORECASE,
        )
        if match:
            return normalize_space(match.group(1))
    match = re.search(
        r"(?:摘要|Abstract)\s*[:：]?\s*(.{40,2000}?)(?=(?:关键词|Keywords|1[.\s、]|引言|Introduction))",
        normalize_space(text),
        flags=re.IGNORECASE,
    )
    return normalize_space(match.group(1)) if match else MISSING


def document_id(relative_path: str) -> str:
    """Create a stable short id from the relative path."""
    return hashlib.sha256(relative_path.encode("utf-8")).hexdigest()[:12]


def ingest_file(
    path: Path, direction: str, project_root: Path = PROJECT_ROOT
) -> dict[str, object]:
    """Extract one supported file into the public index record shape."""
    relative_path = path.relative_to(project_root).as_posix()
    text = READERS[path.suffix.lower()](path)
    cleaned = normalize_space(text)
    warnings: list[str] = []
    if not cleaned:
        warnings.append("未提取到文本；PDF 可能是扫描版或文件内容为空。")
    return {
        "id": document_id(relative_path),
        "file_name": path.name,
        "relative_path": relative_path,
        "file_type": path.suffix.lower().lstrip("."),
        "direction": direction,
        "title": infer_title(text, path),
        "abstract": infer_abstract(text, path),
        "text": cleaned,
        "character_count": len(cleaned),
        "warnings": warnings,
    }


def build_index(root: Path = PROJECT_ROOT) -> dict[str, object]:
    """Scan both literature directions and return an index with error records."""
    documents: list[dict[str, object]] = []
    errors: list[dict[str, str]] = []
    for direction in ("optics", "deep_learning"):
        folder = root / "literature" / direction
        if not folder.exists():
            errors.append(
                {
                    "path": folder.relative_to(root).as_posix(),
                    "direction": direction,
                    "error": "目录不存在。",
                }
            )
            continue
        for path in sorted(folder.rglob("*")):
            if not path.is_file() or path.suffix.lower() not in SUPPORTED_SUFFIXES:
                continue
            try:
                documents.append(ingest_file(path, direction, root))
            except Exception as exc:  # Continue indexing other local files.
                errors.append(
                    {
                        "path": path.relative_to(root).as_posix(),
                        "direction": direction,
                        "error": f"{type(exc).__name__}: {exc}",
                    }
                )
    return {
        "schema_version": 1,
        "generated_at": datetime.now(timezone.utc).isoformat(),
        "project_root": str(root),
        "document_count": len(documents),
        "error_count": len(errors),
        "documents": documents,
        "errors": errors,
    }


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="建立本地文献基础文本索引。")
    parser.add_argument(
        "--output",
        type=Path,
        default=DEFAULT_INDEX_PATH,
        help="索引输出路径，默认 outputs/literature_index.json。",
    )
    return parser.parse_args()


def main() -> int:
    args = parse_args()
    index = build_index()
    args.output.parent.mkdir(parents=True, exist_ok=True)
    args.output.write_text(
        json.dumps(index, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    print(f"索引完成：{index['document_count']} 篇文献，{index['error_count']} 条错误。")
    print(f"输出：{args.output.resolve()}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
